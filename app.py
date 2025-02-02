import streamlit as st
from openai import OpenAI  # Import the OpenAI client

# Initialize the OpenAI client
client = OpenAI(api_key="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")

import streamlit as st
from openai import OpenAI
import PyPDF2
from docx import Document
from io import BytesIO


# Function to extract text from a PDF file
def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    return "".join([page.extract_text() for page in reader.pages])

# Function to extract text from a DOCX file
def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

# Function to interact with OpenAI GPT
def ask_gpt(prompt):
    if prompt in st.session_state:
        return st.session_state[prompt]
    
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "system", "content": "You are a career advisor."}, 
                 {"role": "user", "content": prompt}]
    )
    st.session_state[prompt] = response.choices[0].message.content
    return st.session_state[prompt]

# Function to create a Word document
def create_resume_doc(name, email, phone, skills, experience, education):
    doc = Document()
    doc.add_heading(name, level=1)
    doc.add_paragraph(f"Email: {email} | Phone: {phone}")
    
    doc.add_heading("Skills", level=2)
    for skill in skills.split(","):
        doc.add_paragraph(skill.strip(), style="List Bullet")
    
    doc.add_heading("Experience", level=2)
    doc.add_paragraph(experience)
    
    doc.add_heading("Education", level=2)
    doc.add_paragraph(education)
    
    return doc

# Resume Builder
def resume_builder():
    st.header("Resume Builder")
    
    # Input fields for resume details
    name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone Number")
    skills = st.text_area("Skills (comma-separated)")
    experience = st.text_area("Work Experience")
    education = st.text_area("Education")
    
    if st.button("Build Resume"):
        if not name or not email or not phone or not skills or not experience or not education:
            st.error("Please fill in all fields.")
        else:
            # Create a Word document
            doc = create_resume_doc(name, email, phone, skills, experience, education)
            
            # Save the document to a BytesIO object
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Provide download link
            st.success("Resume built successfully!")
            st.download_button(
                label="Download Resume (Word)",
                data=buffer,
                file_name=f"{name}_Resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Resume Analyzer
def resume_analyzer():
    st.header("AI Career Advisor")
    uploaded_file = st.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])
    
    if uploaded_file:
        # Extract and store resume text
        if 'resume_text' not in st.session_state:
            if uploaded_file.type == "application/pdf":
                st.session_state.resume_text = extract_text_from_pdf(uploaded_file)
            else:
                st.session_state.resume_text = extract_text_from_docx(uploaded_file)
        
        # Suggest Roles button
        if st.button("Suggest Roles"):
            prompt = f"Suggest top 3 career roles for this resume, list only role names:\n{st.session_state.resume_text}"
            roles = ask_gpt(prompt).split("\n")
            st.session_state.roles = [role.strip() for role in roles if role.strip()]
            st.session_state.selected_role = None
            st.session_state.role_data = {}

        # Display role buttons if available
        if 'roles' in st.session_state:
            st.subheader("Suggested Roles")
            cols = st.columns(3)
            for idx, role in enumerate(st.session_state.roles):
                if cols[idx].button(role):
                    st.session_state.selected_role = role
        
        # Show analysis section when role is selected
        if 'selected_role' in st.session_state and st.session_state.selected_role:
            role = st.session_state.selected_role
            st.markdown(f"### Analyzing: {role}")
            
            # Initialize role data structure
            if role not in st.session_state.role_data:
                st.session_state.role_data[role] = {
                    'required_skills': None,
                    'comparison': None,
                    'score': None
                }
            
            # Analysis buttons
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Get Required Skills"):
                    prompt = f"List key technical skills required for {role}, bullet points only"
                    st.session_state.role_data[role]['required_skills'] = ask_gpt(prompt)
            
            with col2:
                if st.button("Run Skill Comparison"):
                    prompt = f"""Compare these resume skills with {role} requirements:
                    Resume: {st.session_state.resume_text}
                    Output format: 
                    - Match Score: X/100 and the calculation of score in table format (required skills, skills in resume, score)
                    - Strong Skills: list
                    - Missing Skills: list with one YouTube playlist and one website for theoretical knowledge
                    - Spelling errors in resume"""
                    st.session_state.role_data[role]['comparison'] = ask_gpt(prompt)
            
            # Display stored data for current role
            if st.session_state.role_data[role]['required_skills']:
                st.subheader("Required Skills")
                st.write(st.session_state.role_data[role]['required_skills'])
            
            if st.session_state.role_data[role]['comparison']:
                st.subheader("Skill Analysis")
                st.write(st.session_state.role_data[role]['comparison'])

# Preparation Guide
def preparation_guide():
    st.header("Preparation Guide")
    if 'selected_role' in st.session_state and st.session_state.selected_role:
        role = st.session_state.selected_role
        st.markdown(f"### Preparing for: {role}")
        
        # Generate a study plan
        if st.button("Generate Study Plan"):
            prompt = f"Create a 30-day study plan for {role}, including daily tasks, resources, and projects."
            study_plan = ask_gpt(prompt)
            st.session_state.study_plan = study_plan
        
        if 'study_plan' in st.session_state:
            st.subheader("Study Plan")
            st.write(st.session_state.study_plan)

# Progress Dashboard
def progress_dashboard():
    st.header("Progress Dashboard")
    if 'test_scores' in st.session_state:
        st.subheader("Test Scores")
        for skill, score in st.session_state.test_scores.items():
            st.write(f"**{skill}:** {score}%")
        
        if 'study_plan' in st.session_state:
            st.subheader("Study Plan Progress")
            st.write(st.session_state.study_plan)

# Main App


def main():
    st.title("AI Career Advisor")
    st.markdown(
        """
        <style>
        .stApp {
            background: linear-gradient(to top, #f971e5, #07017f);
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    st.sidebar.title("Navigation")
    app_mode = st.sidebar.selectbox("Choose a service", ["Resume Builder", "Resume Analyzer","Preparation Guide","Dashboard"])
    
    if app_mode == "Resume Builder":
        resume_builder()
    elif app_mode == "Resume Analyzer":
        resume_analyzer()
    elif app_mode == "Preparation Guide":
        preparation_guide()
    elif app_mode == "Progress Dashboard":
        progress_dashboard()

if __name__ == "__main__":
    main()
